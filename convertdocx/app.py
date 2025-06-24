# app.py

import os
import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from markdown2 import markdown
from bs4 import BeautifulSoup
from config import DOCX_STYLE

INPUT_DIR = "nist-rmf-600-1"
OUTPUT_DIR = "docx"

def apply_style(paragraph, is_heading=False, heading_level=1, is_meta=False):
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    font = run.font
    font.name = DOCX_STYLE["font_name"]
    font.size = Pt(DOCX_STYLE["heading_styles"][f"Heading {heading_level}"]["size"]
                   if is_heading else DOCX_STYLE["font_size"])
    font.bold = DOCX_STYLE["heading_styles"][f"Heading {heading_level}"]["bold"] if is_heading else False

    if is_meta:
        font.color.rgb = RGBColor.from_string(DOCX_STYLE["meta_text_color"])
        paragraph.paragraph_format.line_spacing = 1.0
    elif is_heading:
        font.color.rgb = RGBColor.from_string(DOCX_STYLE["heading_styles"][f"Heading {heading_level}"]["color"])
    else:
        font.color.rgb = RGBColor.from_string(DOCX_STYLE["body_text_color"])
        paragraph.paragraph_format.line_spacing = DOCX_STYLE["line_spacing"]

def convert_markdown_to_docx(md_content, output_path):
    # Extract YAML frontmatter
    frontmatter = {}
    if md_content.startswith("---"):
        parts = md_content.split("---", 2)
        if len(parts) >= 3:
            frontmatter = yaml.safe_load(parts[1])
            md_content = parts[2].strip()

    html = markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(DOCX_STYLE["margin_inches"])
        section.bottom_margin = Inches(DOCX_STYLE["margin_inches"])
        section.left_margin = Inches(DOCX_STYLE["margin_inches"])
        section.right_margin = Inches(DOCX_STYLE["margin_inches"])

    # Render frontmatter
    if frontmatter:
        doc.add_heading("Document Metadata", level=1)
        for key, value in frontmatter.items():
            para = doc.add_paragraph(f"{key}: {value}")
            apply_style(para, is_meta=True)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)


    # Render markdown content
    for elem in soup.descendants:
        if elem.name:
            if elem.name.startswith("h") and elem.name[1:].isdigit():
                level = int(elem.name[1])
                paragraph = doc.add_heading(elem.get_text(), level=level)
                apply_style(paragraph, is_heading=True, heading_level=level)
            elif elem.name == "p":
                paragraph = doc.add_paragraph(elem.get_text())
                apply_style(paragraph)
            elif elem.name == "li":
                bullet = DOCX_STYLE["bullet_style"]["symbol"]
                text = f"{bullet} {elem.get_text()}"
                paragraph = doc.add_paragraph(text)
                apply_style(paragraph)
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)


    doc.save(output_path)

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    for filename in os.listdir(INPUT_DIR):
        if filename.endswith(".md"):
            input_path = os.path.join(INPUT_DIR, filename)
            output_path = os.path.join(OUTPUT_DIR, filename.replace(".md", ".docx"))

            with open(input_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            convert_markdown_to_docx(md_content, output_path)
            print(f"Created: {output_path}")

if __name__ == "__main__":
    main()
